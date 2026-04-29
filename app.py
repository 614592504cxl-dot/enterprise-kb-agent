import os
import io
import re
import json
import time
import sqlite3
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Tuple

import numpy as np
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document

# =========================
# 基础配置
# =========================
load_dotenv()

DB_PATH = os.getenv("KB_DB_PATH", "kb.sqlite3")
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "outputs")
DEFAULT_CHAT_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.2")
DEFAULT_EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# =========================
# 数据库
# =========================
def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    conn = get_conn()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            file_hash TEXT NOT NULL UNIQUE,
            uploaded_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_id INTEGER NOT NULL,
            filename TEXT NOT NULL,
            chunk_index INTEGER NOT NULL,
            text TEXT NOT NULL,
            embedding_json TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(file_id) REFERENCES files(id)
        )
        """
    )
    conn.execute("CREATE INDEX IF NOT EXISTS idx_chunks_file_id ON chunks(file_id)")
    conn.commit()
    conn.close()


# =========================
# OpenAI 客户端
# =========================
@st.cache_resource(show_spinner=False)
def get_openai_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("缺少 OPENAI_API_KEY。请在 .env 文件或系统环境变量中配置。")
    return OpenAI(api_key=api_key)


def embed_texts(texts: List[str], embed_model: str) -> List[List[float]]:
    """批量生成 embedding。"""
    if not texts:
        return []
    client = get_openai_client()
    all_embeddings: List[List[float]] = []
    batch_size = 64
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        response = client.embeddings.create(model=embed_model, input=batch)
        all_embeddings.extend([item.embedding for item in response.data])
    return all_embeddings


def llm_text(prompt: str, instructions: str, chat_model: str) -> str:
    client = get_openai_client()
    response = client.responses.create(
        model=chat_model,
        instructions=instructions,
        input=prompt,
    )
    return response.output_text


# =========================
# 文件解析
# =========================
def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def clean_text(text: str) -> str:
    text = text.replace("\u0000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_upload(uploaded_file) -> Tuple[str, bytes]:
    data = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for idx, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return clean_text("\n\n".join(pages)), data

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 读取表格里的文字
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return clean_text("\n".join(parts)), data

    if filename.endswith(".txt") or filename.endswith(".md"):
        for enc in ["utf-8", "gb18030", "gbk"]:
            try:
                return clean_text(data.decode(enc)), data
            except UnicodeDecodeError:
                pass
        return clean_text(data.decode("utf-8", errors="ignore")), data

    raise ValueError("暂不支持该文件类型。请上传 PDF、DOCX、TXT、MD。")


# =========================
# 文本切块
# =========================
def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 180) -> List[str]:
    """
    轻量切块：优先按段落拼接，超过长度再切。
    chunk_size 不是 token，是字符数。中文场景下先用字符数足够做 MVP。
    """
    text = clean_text(text)
    if not text:
        return []

    paragraphs = re.split(r"\n\s*\n", text)
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(para) > chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            start = 0
            while start < len(para):
                end = start + chunk_size
                piece = para[start:end]
                chunks.append(piece.strip())
                start = max(end - overlap, end)
            continue

        if len(current) + len(para) + 2 <= chunk_size:
            current = f"{current}\n\n{para}" if current else para
        else:
            if current:
                chunks.append(current.strip())
            if overlap > 0 and chunks:
                prev_tail = chunks[-1][-overlap:]
                current = f"{prev_tail}\n\n{para}"
            else:
                current = para

    if current:
        chunks.append(current.strip())

    # 过滤太短的碎片
    return [c for c in chunks if len(c.strip()) >= 20]


# =========================
# 入库
# =========================
def ingest_file(uploaded_file, chunk_size: int, overlap: int, embed_model: str) -> Dict[str, Any]:
    text, raw = extract_text_from_upload(uploaded_file)
    file_hash = sha256_bytes(raw)
    filename = uploaded_file.name

    if not text:
        return {"ok": False, "message": f"{filename} 没有解析出文字。扫描版 PDF 需要 OCR。"}

    conn = get_conn()
    existing = conn.execute("SELECT id FROM files WHERE file_hash = ?", (file_hash,)).fetchone()
    if existing:
        conn.close()
        return {"ok": False, "message": f"{filename} 已经上传过，跳过重复入库。"}

    chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    if not chunks:
        conn.close()
        return {"ok": False, "message": f"{filename} 切块后为空。"}

    embeddings = embed_texts(chunks, embed_model=embed_model)

    cur = conn.execute(
        "INSERT INTO files(filename, file_hash, uploaded_at) VALUES (?, ?, ?)",
        (filename, file_hash, datetime.now().isoformat(timespec="seconds")),
    )
    file_id = cur.lastrowid

    now = datetime.now().isoformat(timespec="seconds")
    rows = [
        (file_id, filename, idx, chunk, json.dumps(embeddings[idx]), now)
        for idx, chunk in enumerate(chunks)
    ]
    conn.executemany(
        """
        INSERT INTO chunks(file_id, filename, chunk_index, text, embedding_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        rows,
    )
    conn.commit()
    conn.close()

    return {"ok": True, "message": f"{filename} 入库完成：{len(chunks)} 个文本块。"}


# =========================
# 检索
# =========================
def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def retrieve(query: str, top_k: int, embed_model: str) -> List[Dict[str, Any]]:
    q_vec = np.array(embed_texts([query], embed_model=embed_model)[0], dtype=np.float32)

    conn = get_conn()
    rows = conn.execute(
        "SELECT id, filename, chunk_index, text, embedding_json FROM chunks"
    ).fetchall()
    conn.close()

    scored: List[Dict[str, Any]] = []
    for row in rows:
        try:
            vec = np.array(json.loads(row["embedding_json"]), dtype=np.float32)
            score = cosine_similarity(q_vec, vec)
            scored.append({
                "id": row["id"],
                "filename": row["filename"],
                "chunk_index": row["chunk_index"],
                "text": row["text"],
                "score": score,
            })
        except Exception:
            continue

    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]


def build_context(docs: List[Dict[str, Any]], max_chars: int = 12000) -> str:
    parts = []
    total = 0
    for idx, d in enumerate(docs, start=1):
        text = d["text"].strip()
        block = (
            f"[来源 {idx}]\n"
            f"文件：{d['filename']}\n"
            f"文本块：{d['chunk_index']}\n"
            f"相关度：{d['score']:.4f}\n"
            f"内容：\n{text}\n"
        )
        if total + len(block) > max_chars:
            break
        parts.append(block)
        total += len(block)
    return "\n\n".join(parts)


# =========================
# Agent 逻辑
# =========================
def answer_agent(question: str, retrieved_docs: List[Dict[str, Any]], chat_model: str) -> str:
    context = build_context(retrieved_docs)
    instructions = """
你是一个企业内部知识库问答 Agent。你的规则：
1. 只基于提供的资料回答，不要编造。
2. 如果资料不足，明确说“资料中没有找到足够依据”。
3. 回答必须带来源编号，例如 [来源 1]、[来源 2]。
4. 输出要简洁、结构化，适合企业员工直接使用。
5. 如问题涉及流程、制度、合同、客户历史，优先提取关键事实、责任人、时间、条件、风险。
""".strip()

    prompt = f"""
用户问题：
{question}

可用资料：
{context}

请按这个格式回答：
## 结论

## 依据

## 风险 / 不确定点

## 下一步建议
""".strip()
    return llm_text(prompt, instructions=instructions, chat_model=chat_model)


def action_agent(objective: str, retrieved_docs: List[Dict[str, Any]], action_type: str, chat_model: str) -> str:
    context = build_context(retrieved_docs)
    instructions = """
你是一个企业自动办事 Agent。你不能真的发送邮件、修改系统或替用户做不可逆操作。
你只能基于资料生成：待办清单、邮件草稿、会议纪要、项目推进表、风险清单、执行步骤。
输出要可复制、可执行、清楚标注依据来源。
""".strip()

    prompt = f"""
用户想完成的事项：
{objective}

办事类型：
{action_type}

可用资料：
{context}

请生成一个可执行结果。要求：
1. 不编造资料中没有的信息；
2. 缺失信息单独列出；
3. 需要人工确认的地方打上【需确认】；
4. 引用来源编号，例如 [来源 1]；
5. 根据办事类型输出对应格式。
""".strip()
    return llm_text(prompt, instructions=instructions, chat_model=chat_model)


def save_markdown(title: str, content: str) -> str:
    safe_title = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", title).strip("_")[:40] or "agent_output"
    filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{safe_title}.md"
    path = os.path.join(OUTPUT_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


# =========================
# UI
# =========================
def render_source_cards(docs: List[Dict[str, Any]]) -> None:
    if not docs:
        return
    with st.expander("查看本次引用的知识库片段", expanded=False):
        for idx, d in enumerate(docs, start=1):
            st.markdown(f"**来源 {idx}｜{d['filename']}｜文本块 {d['chunk_index']}｜相关度 {d['score']:.4f}**")
            st.write(d["text"][:1200] + ("..." if len(d["text"]) > 1200 else ""))
            st.divider()


def render_sidebar() -> Dict[str, Any]:
    st.sidebar.title("配置")
    chat_model = st.sidebar.text_input("问答模型", value=DEFAULT_CHAT_MODEL)
    embed_model = st.sidebar.text_input("Embedding 模型", value=DEFAULT_EMBED_MODEL)
    top_k = st.sidebar.slider("检索片段数", min_value=3, max_value=15, value=6)
    chunk_size = st.sidebar.slider("切块字符数", min_value=500, max_value=2500, value=1200, step=100)
    overlap = st.sidebar.slider("切块重叠字符数", min_value=0, max_value=500, value=180, step=20)

    st.sidebar.divider()
    conn = get_conn()
    file_count = conn.execute("SELECT COUNT(*) AS c FROM files").fetchone()["c"]
    chunk_count = conn.execute("SELECT COUNT(*) AS c FROM chunks").fetchone()["c"]
    conn.close()
    st.sidebar.metric("文件数", file_count)
    st.sidebar.metric("文本块", chunk_count)

    return {
        "chat_model": chat_model.strip(),
        "embed_model": embed_model.strip(),
        "top_k": top_k,
        "chunk_size": chunk_size,
        "overlap": overlap,
    }


def page_upload(config: Dict[str, Any]) -> None:
    st.header("1. 上传企业资料")
    st.caption("支持 PDF、DOCX、TXT、MD。扫描版 PDF 暂不做 OCR。")

    files = st.file_uploader(
        "上传文件",
        type=["pdf", "docx", "txt", "md"],
        accept_multiple_files=True,
    )

    if st.button("开始入库", type="primary", disabled=not files):
        for f in files:
            try:
                with st.spinner(f"正在处理 {f.name} ..."):
                    result = ingest_file(
                        f,
                        chunk_size=config["chunk_size"],
                        overlap=config["overlap"],
                        embed_model=config["embed_model"],
                    )
                if result["ok"]:
                    st.success(result["message"])
                else:
                    st.warning(result["message"])
            except Exception as e:
                st.error(f"{f.name} 处理失败：{e}")

    st.subheader("已入库文件")
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, filename, uploaded_at FROM files ORDER BY id DESC LIMIT 50"
    ).fetchall()
    conn.close()
    if rows:
        st.dataframe([dict(r) for r in rows], use_container_width=True)
    else:
        st.info("还没有入库文件。")


def page_qa(config: Dict[str, Any]) -> None:
    st.header("2. 知识库问答")
    question = st.text_area(
        "输入问题",
        placeholder="例如：报销流程是什么？这个客户之前有哪些关键沟通？某项制度里对审批有什么要求？",
        height=120,
    )

    if st.button("检索并回答", type="primary", disabled=not question.strip()):
        try:
            with st.spinner("正在检索知识库..."):
                docs = retrieve(question, top_k=config["top_k"], embed_model=config["embed_model"])
            if not docs:
                st.warning("知识库为空或没有检索到内容。")
                return
            with st.spinner("正在生成回答..."):
                answer = answer_agent(question, docs, chat_model=config["chat_model"])
            st.markdown(answer)
            render_source_cards(docs)
            path = save_markdown(question, answer)
            st.caption(f"已保存到：{path}")
        except Exception as e:
            st.error(f"问答失败：{e}")


def page_action(config: Dict[str, Any]) -> None:
    st.header("3. 自动办事")
    action_type = st.selectbox(
        "办事类型",
        ["待办清单", "邮件草稿", "会议纪要", "项目推进表", "风险清单", "执行步骤"],
    )
    objective = st.text_area(
        "你要让 Agent 做什么",
        placeholder="例如：根据客户沟通记录，帮我写一封跟进邮件；根据会议纪要整理本周项目推进清单。",
        height=140,
    )

    if st.button("生成办事结果", type="primary", disabled=not objective.strip()):
        try:
            with st.spinner("正在检索相关资料..."):
                docs = retrieve(objective, top_k=config["top_k"], embed_model=config["embed_model"])
            if not docs:
                st.warning("知识库为空或没有检索到内容。")
                return
            with st.spinner("正在生成办事结果..."):
                result = action_agent(objective, docs, action_type=action_type, chat_model=config["chat_model"])
            st.markdown(result)
            render_source_cards(docs)
            path = save_markdown(objective, result)
            st.caption(f"已保存到：{path}")
        except Exception as e:
            st.error(f"自动办事失败：{e}")


def page_admin() -> None:
    st.header("4. 管理")
    st.warning("删除知识库会清空本地 SQLite 数据库，文件需要重新上传入库。")
    if st.button("清空知识库", type="secondary"):
        conn = get_conn()
        conn.execute("DELETE FROM chunks")
        conn.execute("DELETE FROM files")
        conn.commit()
        conn.close()
        st.success("已清空。请刷新页面。")

    st.subheader("数据库位置")
    st.code(os.path.abspath(DB_PATH))
    st.subheader("输出文件夹")
    st.code(os.path.abspath(OUTPUT_DIR))


def main() -> None:
    st.set_page_config(page_title="企业知识库问答 + 自动办事 Agent", layout="wide")
    init_db()

    st.title("企业知识库问答 + 自动办事 Agent")
    st.caption("MVP：本地 SQLite 向量库 + OpenAI Embeddings + Responses API + Streamlit 界面")

    config = render_sidebar()

    tab1, tab2, tab3, tab4 = st.tabs(["上传入库", "知识库问答", "自动办事", "管理"])
    with tab1:
        page_upload(config)
    with tab2:
        page_qa(config)
    with tab3:
        page_action(config)
    with tab4:
        page_admin()


if __name__ == "__main__":
    main()
